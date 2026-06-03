"""Word .docx reader via python-docx."""

from pathlib import Path

MAX_CHARS = 50_000


def read(path: Path) -> str:
    from docx import Document
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    out = "\n".join(parts)
    if len(out) > MAX_CHARS:
        out = out[:MAX_CHARS] + f"\n\n... (truncated at {MAX_CHARS} chars)"
    return out or "(empty document)"
